#!/usr/bin/env python3
"""
AI Talent Engine — Interview Demo Document Generator
Version: Day4-v1.1 (Hardened)
Date: 2025-12-30
Author: L. David Mendoza
© 2025 L. David Mendoza. All Rights Reserved.

Purpose
-------
Generate all interview-ready Word documents in a single, deterministic run.
This script has NO side effects beyond writing .docx files to the repo root.

Hardening
---------
Fails fast with a self-explanatory error if python-docx is not installed.
"""

# -------------------------
# Dependency Guard (Explicit)
# -------------------------
try:
    from docx import Document
except ImportError:
    raise SystemExit(
        "ERROR: python-docx is not installed.\n\n"
        "This script generates Word (.docx) files and requires python-docx.\n\n"
        "Install with:\n"
        "  python3 -m pip install --user python-docx\n"
    )

# -------------------------
# Helpers
# -------------------------
def write_doc(path: str, title: str, body_lines: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body_lines:
        doc.add_paragraph(line)
    doc.save(path)

# -------------------------
# Main Generator
# -------------------------
def main() -> None:

    # 1. Interview Demo Script (Read-Aloud)
    write_doc(
        "INTERVIEW_DEMO_SCRIPT.docx",
        "Interview Demo Script — Read-Aloud Version",
        [
            "Version: Day4-v1.1",
            "Author: L. David Mendoza",
            "© 2025 L. David Mendoza. All Rights Reserved.",
            "",
            "Opening (Always Say This First)",
            "“Before I run anything, I’ll explain what question this scenario answers "
            "and why this execution envelope is appropriate.”",
            "",
            "Scenario Declaration",
            "“I’m running the scenario: demo: <SCENARIO NAME>.”",
            "“The demo envelope is intentionally bounded for clarity and speed.”",
            "“The exhaustive version uses the same logic and runs asynchronously for completeness.”",
            "",
            "Intent Statement",
            "“This scenario exists to answer <PRIMARY QUESTION>.”",
            "“It intentionally does not attempt to answer <EXCLUDED QUESTION>.”",
            "",
            "Boundary Condition",
            "“If this scenario were used for <WRONG USE CASE>, "
            "the results would be misleading — and that’s by design.”",
            "",
            "Governance Statement",
            "“Nothing in this system runs unless it’s explicitly allowed, declared, and auditable.”",
            "",
            "Production Close",
            "“In production, I run this scenario exhaustively because completeness matters more "
            "than speed once correctness is established.”",
        ],
    )

    # 2. Interview Cheat Sheet (One Page)
    write_doc(
        "INTERVIEW_DEMO_CHEAT_SHEET.docx",
        "Interview Demo Cheat Sheet (One Page)",
        [
            "Version: Day4-v1.1",
            "",
            "Before Running Anything",
            "- Name the scenario exactly",
            "- Declare demo vs exhaustive",
            "- State what it does NOT do",
            "",
            "Always Say",
            "- This is a demo envelope, bounded for clarity.",
            "- The logic never changes, only the execution envelope.",
            "- No single signal is determinative.",
            "- Disagreement does not require re-execution.",
            "",
            "If Asked About Runtime",
            "- Demo runs are intentionally short",
            "- Exhaustive runs execute asynchronously",
            "",
            "Never Do Live",
            "- Improvise prompts",
            "- Change logic mid-demo",
            "- Apologize for depth",
            "",
            "Close With",
            "“Once correctness is established, I flip to exhaustive mode for completeness.”",
        ],
    )

    # 3. Interview Red-Team Drill (Hostile Questions)
    write_doc(
        "INTERVIEW_RED_TEAM_DRILL.docx",
        "Interview Red-Team Drill — Hostile Questioning",
        [
            "Version: Day4-v1.1",
            "",
            "Why this scenario?",
            "Answer: It isolates the correct signal while explicitly excluding noise.",
            "",
            "When would this fail?",
            "Answer: When used outside its declared scope — intentionally.",
            "",
            "Why not just use an LLM?",
            "Answer: Because evidence must survive model changes.",
            "",
            "Does this scale?",
            "Answer: The logic is invariant; only execution cost scales.",
            "",
            "What did you choose not to automate?",
            "Answer: Final human judgment remains intentional.",
        ],
    )

    print("Interview demo documents generated successfully.")

# -------------------------
# Entry Point
# -------------------------
if __name__ == "__main__":
    main()
